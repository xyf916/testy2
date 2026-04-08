#!/usr/bin/env python3
"""
FR Document Ingestion CLI

Recursively scans configured folders, extracts text and metadata from
.docx and .doc files, and populates the SQLite database with full-text
indexes for fast keyword search.

Usage:
    python ingest_fr.py <db_path> <folder1> [folder2 ...] [--include PATTERN] [--exclude PATTERN]

Progress is streamed as JSON lines to stdout.
"""

import sys
import os
import json
import re
import sqlite3
import time
import argparse
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import win32com.client  # type: ignore[import]
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False


# ---------------------------------------------------------------------------
# Database setup
# ---------------------------------------------------------------------------

def init_search_db(db_path):
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    conn = sqlite3.connect(db_path)
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS fr_documents (
            fr_number   TEXT PRIMARY KEY,
            title       TEXT,
            originator  TEXT,
            file_path   TEXT,
            ingested_at INTEGER
        );

        CREATE VIRTUAL TABLE IF NOT EXISTS fr_fts USING fts5(
            fr_number UNINDEXED,
            title,
            body_text,
            tokenize='porter unicode61'
        );
    ''')
    conn.commit()
    conn.close()


def upsert_document(conn, fr_number, title, originator, file_path, body_text):
    conn.execute("DELETE FROM fr_fts WHERE fr_number = ?", (fr_number,))
    conn.execute(
        'INSERT OR REPLACE INTO fr_documents (fr_number, title, originator, file_path, ingested_at) VALUES (?, ?, ?, ?, ?)',
        (fr_number, title, originator, file_path, int(time.time())),
    )
    conn.execute(
        "INSERT INTO fr_fts (fr_number, title, body_text) VALUES (?, ?, ?)",
        (fr_number, title, body_text),
    )


# ---------------------------------------------------------------------------
# Metadata extraction (best-effort)
# ---------------------------------------------------------------------------

def extract_originator(text):
    m = re.search(r'ORIGINATOR\s*\n\s*(.+)', text, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return ""


# ---------------------------------------------------------------------------
# Document readers
# ---------------------------------------------------------------------------

def read_docx_for_ingest(file_path):
    """Return (title, body_text) from a .docx file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(file_path)
    parts = []
    title = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if not title and ("heading" in style or "title" in style):
            title = text
        parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                t = cell.text.strip()
                if t:
                    parts.append(t)

    body = "\n".join(parts)
    if not title and parts:
        title = parts[0]

    return title, body


def doc_to_docx(file_path):
    """Convert a .doc file to a temp .docx using Word, return the temp path."""
    import pythoncom  # type: ignore[import]
    import tempfile
    import shutil

    tmp_dir = tempfile.gettempdir()
    pid = os.getpid()
    tmp_doc  = os.path.join(tmp_dir, f"fr_ingest_tmp_{pid}.doc")
    tmp_docx = os.path.join(tmp_dir, f"fr_ingest_tmp_{pid}.docx")
    shutil.copy2(file_path, tmp_doc)

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(tmp_doc))
        doc.SaveAs(tmp_docx, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
        doc.Close(0)
    finally:
        try:
            if word:
                word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()
        try:
            os.remove(tmp_doc)
        except Exception:
            pass
    return tmp_docx


def read_doc_for_ingest(file_path):
    if not WIN32_AVAILABLE:
        raise ValueError("pywin32 not installed. Run: pip install pywin32")
    tmp = doc_to_docx(file_path)
    try:
        return read_docx_for_ingest(tmp)
    finally:
        try:
            os.remove(tmp)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# File collection
# ---------------------------------------------------------------------------

def extract_fr_number(file_path):
    """Pull the numeric FR ID out of a filename, e.g. FR222.docx → '222'."""
    name = os.path.basename(file_path)
    m = re.search(r'FR(\d+)', name, re.IGNORECASE)
    return m.group(1) if m else None


def is_excluded(file_path, exclude_patterns):
    if not exclude_patterns:
        return False
    p = Path(file_path)
    for pat in exclude_patterns:
        try:
            if p.match(pat):
                return True
        except Exception:
            pass
    return False


def matches_include(file_path, include_patterns):
    if not include_patterns:
        return True
    p = Path(file_path)
    for pat in include_patterns:
        try:
            if p.match(pat):
                return True
        except Exception:
            pass
    return False


def collect_files(root_folders, include_patterns, exclude_patterns):
    """Recursively collect all .docx/.doc files matching the given patterns."""
    files = []
    for folder in root_folders:
        folder = os.path.normpath(folder)
        if not os.path.isdir(folder):
            emit({"stage": "warning", "message": f"Folder not found, skipping: {folder}"})
            continue
        for dirpath, _, filenames in os.walk(folder):
            for filename in filenames:
                ext = os.path.splitext(filename)[1].lower()
                if ext not in ('.docx', '.doc'):
                    continue
                full_path = os.path.join(dirpath, filename)
                if not matches_include(full_path, include_patterns):
                    continue
                if is_excluded(full_path, exclude_patterns):
                    continue
                files.append(full_path)
    return files


# ---------------------------------------------------------------------------
# Ingestion loop
# ---------------------------------------------------------------------------

def run_ingestion(db_path, root_folders, include_patterns, exclude_patterns):
    init_search_db(db_path)

    emit({"stage": "scanning", "message": "Scanning folders..."})
    files = collect_files(root_folders, include_patterns, exclude_patterns)
    total = len(files)
    emit({"stage": "found", "total": total})

    if total == 0:
        emit({"stage": "done", "count": 0, "errors": 0,
              "message": "No .docx or .doc files found in the configured folders."})
        return

    conn = sqlite3.connect(db_path)
    count = 0
    errors = 0

    for i, file_path in enumerate(files):
        fr_number = extract_fr_number(file_path)
        if not fr_number:
            emit({
                "stage": "skip",
                "file": os.path.basename(file_path),
                "reason": "No FR number found in filename",
            })
            continue

        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".docx":
                title, body = read_docx_for_ingest(file_path)
            elif ext == ".doc":
                title, body = read_doc_for_ingest(file_path)
            else:
                continue

            originator = extract_originator(body)
            upsert_document(conn, fr_number, title, originator, file_path, body)
            count += 1

            # Commit in batches to avoid holding a huge transaction
            if count % 50 == 0:
                conn.commit()

            emit({
                "stage": "progress",
                "current": i + 1,
                "total": total,
                "fr_number": fr_number,
                "file": os.path.basename(file_path),
            })

        except Exception as e:
            errors += 1
            emit({
                "stage": "error",
                "file": os.path.basename(file_path),
                "error": str(e),
            })

    conn.commit()
    conn.close()

    emit({"stage": "done", "count": count, "errors": errors})


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def emit(obj):
    """Write a JSON line to stdout immediately."""
    print(json.dumps(obj), flush=True)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Ingest FR documents into the SQLite search database."
    )
    parser.add_argument("db_path", help="Path to the SQLite database file")
    parser.add_argument("folders", nargs="+", help="Root folders to scan")
    parser.add_argument(
        "--include",
        action="append",
        dest="include_patterns",
        default=None,
        metavar="PATTERN",
        help="Glob include pattern (default: **/*.docx and **/*.doc)",
    )
    parser.add_argument(
        "--exclude",
        action="append",
        dest="exclude_patterns",
        default=None,
        metavar="PATTERN",
        help="Glob exclude pattern",
    )

    args = parser.parse_args()
    include_patterns = args.include_patterns or ["**/*.docx", "**/*.doc"]
    exclude_patterns = args.exclude_patterns or []

    run_ingestion(args.db_path, args.folders, include_patterns, exclude_patterns)


if __name__ == "__main__":
    main()
